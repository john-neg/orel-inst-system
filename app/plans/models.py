import os

import requests
from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter

from app.common.func.app_core import xlsx_iter_rows, xlsx_normalize
from app.common.reports.ExcelStyles import ExcelStyle
from app.main.func import db_filter_req, db_request
from app.main.models import EducationPlan
from app.plans.func import (
    disciplines_wp_clean,
    disciplines_comp_load,
    disciplines_comp_del,
)
from config import FlaskConfig, ApeksConfig as Apeks


class CompPlan(EducationPlan):
    def __init__(self, education_plan_id):
        super().__init__(education_plan_id)
        self.competencies = self.get_comp()

    def get_comp(self):
        """Получение списка компетенций плана"""
        return db_filter_req(
            "plan_competencies", "education_plan_id", self.education_plan_id
        )

    def get_comp_list(self):
        comp_list = []
        for comp in self.competencies:
            comp_list.append(comp.get("code"))
        return comp_list

    def get_comp_code_by_id(self, comp_id):
        """Получение кода компетенции по id"""
        for comp in self.competencies:
            if comp.get("id") == comp_id:
                return comp.get("code")

    def get_comp_by_id(self, comp_id):
        """Получение кода и названия компетенции по id"""
        for comp in self.competencies:
            if comp.get("id") == comp_id:
                return f'{comp.get("code")} - {comp.get("description")}'

    def get_comp_id_by_code(self, comp_code):
        """Получение id компетенции по коду"""
        for comp in self.competencies:
            if comp.get("code") == comp_code:
                return comp.get("id")

    def load_comp(self, code, description, left_node, right_node):
        """Добавление компетенции и места в списке"""
        params = {"token": Apeks.TOKEN}
        data = {
            "table": "plan_competencies",
            "fields[education_plan_id]": self.education_plan_id,
            "fields[code]": code,
            "fields[description]": description,
            "fields[level]": "1",
            "fields[left_node]": str(left_node),
            "fields[right_node]": str(right_node),
        }
        load = requests.post(
            Apeks.URL + "/api/call/system-database/add", params=params, data=data
        )
        if load.json()["status"] == 0:
            return f"{code} {description} {load.json()['message']}"
        self.competencies = self.get_comp()

    def del_comp(self):
        """
        Удаляет все компетенции из плана
        (если нельзя (связаны), то сообщение)
        """
        message = []
        for i in range(len(self.competencies)):
            params = {
                "token": Apeks.TOKEN,
                "table": "plan_competencies",
                "filter[id]": self.competencies[i]["id"],
            }
            remove = requests.delete(
                Apeks.URL + "/api/call/system-database/delete", params=params
            )
            if remove.json()["status"] == 0:
                message.append(
                    f"{self.competencies[i]['code']} " + f"{remove.json()['message']}"
                )
        self.competencies = self.get_comp()
        return message

    def disciplines_comp(self):
        """Получение связей дисциплин и компетенций"""
        message = []
        for disc in self.disciplines.keys():
            params = {
                "token": Apeks.TOKEN,
                "table": "plan_curriculum_discipline_competencies",
                "filter[curriculum_discipline_id]": disc,
            }
            resp = requests.get(
                Apeks.URL + "/api/call/system-database/get", params=params
            )
            if resp.json()["data"]:
                message += resp.json()["data"]
        return message

    def disciplines_comp_dict(self):
        """Получение связей дисциплин и компетенций с названиями"""
        data = self.disciplines_comp()
        mtrx_dict = {}
        disc_id = 0
        if data:
            for i in data:
                if disc_id == i["curriculum_discipline_id"]:
                    mtrx_dict[
                        self.discipline_name(i["curriculum_discipline_id"])
                    ].append(self.get_comp_by_id(i["competency_id"]))
                else:
                    mtrx_dict[self.discipline_name(i["curriculum_discipline_id"])] = [
                        self.get_comp_by_id(i["competency_id"])
                    ]
                    disc_id = i["curriculum_discipline_id"]
        return mtrx_dict

    def disciplines_all_comp_del(self):
        """
        Удаление компетенций, всех связей с дисциплинами и содержания из РП.
        """
        for curriculum_discipline_id in self.disciplines.keys():
            work_program_list = db_filter_req(
                "mm_work_programs", "curriculum_discipline_id", curriculum_discipline_id
            )
            for wp in work_program_list:
                if wp.get("id"):
                    disciplines_wp_clean(wp.get("id"))
            disciplines_comp_del(curriculum_discipline_id)
            self.del_comp()

    def matrix_generate(self):
        """Формирование матрицы компетенций плана в формате Excel"""
        disc_list = db_request("plan_disciplines")
        plan_data = db_filter_req(
            "plan_curriculum_disciplines", "education_plan_id", self.education_plan_id
        )
        relations = self.disciplines_comp()

        # Сортировка
        plan_disc = {}
        for disc in plan_data:
            plan_disc[int(disc["left_node"])] = {
                "id": disc["id"],
                "code": disc["code"],
                "discipline_id": disc["discipline_id"],
                "level": disc["level"],
            }
        plan_comp = {}
        for comp in self.competencies:
            plan_comp[int(comp["left_node"])] = {"id": comp["id"], "code": comp["code"]}

        def disc_name(discipline_id):
            for discipline in disc_list:
                if discipline.get("id") == discipline_id:
                    return discipline["name"]

        wb = Workbook()
        ws = wb.active
        ws.title = "Матрица компетенций"
        ws.cell(1, 1).value = "Код"
        ws.cell(1, 1).style = ExcelStyle.Header
        ws.column_dimensions[get_column_letter(1)].width = 15
        ws.cell(1, 2).value = "Название дисциплины"
        ws.cell(1, 2).style = ExcelStyle.Header
        ws.column_dimensions[get_column_letter(2)].width = 60

        row = 2
        for disc in sorted(plan_disc):
            ws.cell(row, 1).value = plan_disc[disc]["code"]
            ws.cell(row, 1).style = ExcelStyle.Base
            ws.cell(row, 2).value = disc_name(plan_disc[disc]["discipline_id"])
            ws.cell(row, 2).style = ExcelStyle.Base
            column = 3
            for comp in sorted(plan_comp):
                ws.cell(1, column).value = plan_comp[comp]["code"]
                ws.cell(1, column).style = ExcelStyle.Header
                ws.cell(1, column).alignment = Alignment(text_rotation=90)
                ws.column_dimensions[get_column_letter(column)].width = 4
                ws.cell(row, column).style = ExcelStyle.Base
                ws.cell(row, column).alignment = Alignment(
                    horizontal="center", vertical="center"
                )
                if plan_disc[disc]["level"] != "3":
                    ws.cell(row, 1).style = ExcelStyle.BaseBold
                    ws.cell(row, 1).fill = ExcelStyle.GreyFill
                    ws.cell(row, 2).style = ExcelStyle.BaseBold
                    ws.cell(row, 2).fill = ExcelStyle.GreyFill
                    ws.cell(row, column).style = ExcelStyle.BaseBold
                    ws.cell(row, column).fill = ExcelStyle.GreyFill
                for relation in relations:
                    if (
                        plan_disc[disc]["id"] == relation["curriculum_discipline_id"]
                        and plan_comp[comp]["id"] == relation["competency_id"]
                    ):
                        ws.cell(row, column).value = "+"
                column += 1
            row += 1

        filename = f"Матрица - {self.name}.xlsx"
        wb.save(FlaskConfig.EXPORT_FILE_DIR + filename)
        return filename

    def matrix_delete(self):
        """Удаление связей дисциплин и компетенций и их содержания в РП"""
        for curriculum_discipline_id in self.disciplines.keys():
            work_program_list = db_filter_req(
                "mm_work_programs", "curriculum_discipline_id", curriculum_discipline_id
            )
            for wp in work_program_list:
                if wp.get("id"):
                    disciplines_wp_clean(wp.get("id"))
            disciplines_comp_del(curriculum_discipline_id)

    def matrix_simple_file_check(self, filename):
        """Проверка файла простой матрицы"""
        wb = load_workbook(filename)
        ws = wb.active
        file_data = list(xlsx_iter_rows(ws))
        report = {}
        comp_code_errors = []
        for disc in self.disciplines:
            report[self.disciplines[disc][1]] = []
            for line in file_data:
                if self.disciplines[disc][1] == line[1]:
                    for i in range(2, len(line)):
                        if str(line[i]) == "+":
                            report[self.disciplines[disc][1]].append(file_data[0][i])
                    if not report[self.disciplines[disc][1]]:
                        report[self.disciplines[disc][1]] = ["None"]
        for i in range(2, len(file_data[0])):
            if not self.get_comp_id_by_code(file_data[0][i]):
                comp_code_errors.append(file_data[0][i])
        return report, comp_code_errors

    def matrix_simple_upload(self, filename):
        """Загрузка связей из файла простой матрицы"""
        wb = load_workbook(filename)
        ws = wb.active
        file_data = list(xlsx_iter_rows(ws))
        load_data = {}
        for disc in self.disciplines:
            load_data[disc] = []
            for line in file_data:
                if self.disciplines[disc][1] == line[1]:
                    for i in range(2, len(line)):
                        if str(line[i]) == "+":
                            load_data[disc].append(
                                self.get_comp_id_by_code(file_data[0][i])
                            )
        for data in load_data:
            if load_data.get(data):
                for comp in load_data.get(data):
                    if comp is not None:
                        disciplines_comp_load(data, comp)


class MatrixIndicatorsFile:
    def __init__(self, filename):
        self.wb = load_workbook(filename)
        self.ws = self.wb.active
        self.normalize()
        self.file_data = list(xlsx_iter_rows(self.ws))
        self.disciplines_list = self.get_disciplines_list()
        self.title_name = os.path.splitext(os.path.basename(filename))[0]

    def normalize(self):
        """Коррекция ошибок в загруженном файле"""
        replace_dict = {
            "     ": " ",
            "    ": " ",
            "   ": " ",
            "  ": " ",
            "–": "-",
            ". - ": " - ",
            "K": "К",  # Eng to RUS
            "O": "О",
            "A": "А",
            "B": "В",
            "C": "С",
            "H": "Н",
            "y": "у",
            ". з.": ".з.",
            ". у.": ".у.",
            ". в.": ".в.",
            ".з .": ".з.",
            ".у .": ".у.",
            ".в .": ".в.",
            "None": "",
        }
        xlsx_normalize(self.ws, replace_dict)

    def file_errors(self):
        """
        Проверка матрицы на ошибки распознавания индикаторов
        (знать, уметь, владеть).
        """
        not_found = []
        ext_to_check = [".з.", ".у.", ".в."]
        for k in range(2, len(self.file_data[0])):
            if self.file_data[0][k] and self.file_data[0][k] != "None":
                if all(ext not in self.file_data[0][k] for ext in ext_to_check):
                    not_found.append(self.file_data[0][k])
                if " - " not in self.file_data[0][k]:
                    not_found.append(
                        'Проверьте разделитель " - " ' + self.file_data[0][k]
                    )
        return not_found

    def get_disciplines_list(self):
        """Получение списка дисциплин загруженного файла."""
        disc_list = []
        for row in range(1, len(self.file_data)):
            if self.file_data[row][1] and self.file_data[row][1] != "None":
                disc_list.append(self.file_data[row][1])
        # for row in range(2, self.ws.max_row):
        #     if self.ws.cell(row, 2).value != 'None':
        #         disc_list.append(self.ws.cell(row, 2).value)
        return disc_list

    def find_discipline(self, discipline_name):
        """
        Проверяем есть ли дисциплина в загруженной матрице (Google Sheets).
        """
        return True if discipline_name in self.disciplines_list else False

    def disc_comp(self, discipline_name):
        """Компетенция и ее индикаторы для дисциплины."""
        comp, disc_comp = "", {}
        for row in range(1, len(self.disciplines_list) + 1):
            if self.file_data[row][1]:
                if discipline_name.strip() == self.file_data[row][1]:
                    for col in range(len(self.file_data[0])):
                        if self.file_data[row][col] == "+":
                            # Получаем компетенцию
                            if comp != self.file_data[0][col].split(".")[0]:
                                comp = self.file_data[0][col].split(".")[0]
                                disc_comp[comp] = {
                                    "knowledge": [],
                                    "abilities": [],
                                    "ownerships": [],
                                }
                            # Получаем индикатор
                            data = self.file_data[0][col]
                            load_data = (
                                f"{data.split(' - ')[1]} " f"({data.split(' - ')[0]})"
                            )
                            load_data = load_data.replace("  ", " ").replace(
                                ". (", " ("
                            )
                            if self.file_data[row][col] == "+" and ".з." in str(data):
                                disc_comp[comp]["knowledge"] += [load_data]
                            elif self.file_data[row][col] == "+" and ".у." in str(data):
                                disc_comp[comp]["abilities"] += [load_data]
                            elif self.file_data[row][col] == "+" and ".в." in str(data):
                                disc_comp[comp]["ownerships"] += [load_data]
        return disc_comp

    def comp_list(self):
        comp, comp_list = "", []
        """Список компетенций присутствующих в матрице"""
        for col in range(2, len(self.file_data[0])):
            if comp != self.file_data[0][col].split(".")[0]:
                comp = self.file_data[0][col].split(".")[0]
                comp_list.append(comp)
        return comp_list

    def list_to_word(self):
        """ "Конвертация файла xlsx в Word (Индикаторы по дисциплинам)"""
        document = Document()
        section = document.sections[-1]
        section.top_margin = Cm(1)  # Верхний отступ
        section.bottom_margin = Cm(1)  # Нижний отступ
        section.left_margin = Cm(1)  # Отступ слева
        section.right_margin = Cm(1)  # Отступ справа
        paragraph_format = document.styles["Normal"].paragraph_format
        paragraph_format.line_spacing_rule = (
            WD_LINE_SPACING.SINGLE
        )  # межстрочный интервал
        paragraph_format.space_after = Pt(0)  # между абзацами
        style = document.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"  # Стиль шрифта
        font.size = Pt(12)
        document.add_heading(self.title_name, 0)  # Заголовок
        document.add_paragraph("Индикаторы компетенций по дисциплинам")
        for row in range(1, len(self.disciplines_list) + 1):
            document.add_page_break()
            document.add_heading(self.file_data[row][1], level=1)
            # Список компетенций
            disc_data = self.disc_comp(self.file_data[row][1])
            document.add_paragraph(", ".join(list(disc_data.keys())))
            # Индикаторы
            indicator_types = {
                "knowledge": "Знать:",
                "abilities": "Уметь:",
                "ownerships": "Владеть:",
            }
            for typ in indicator_types:
                document.add_paragraph("")
                p = document.add_paragraph("")
                p.add_run(indicator_types[typ]).bold = True
                for d in disc_data:
                    for ind in range(len(disc_data[d][typ])):
                        paragraph = document.add_paragraph(
                            "- " + disc_data[d][typ][ind]
                        )
                        paragraph.add_run(";")
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        document.save(FlaskConfig.EXPORT_FILE_DIR + self.title_name + ".docx")
        return self.title_name + ".docx"
