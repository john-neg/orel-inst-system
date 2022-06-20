import logging

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt

from config import FlaskConfig


def generate_program_title_pages(form_data: dict, plan_name: str, wp_data: dict) -> str:
    """
    Формирует титульные листы рабочих программ учебного плана в формате docx.

    Parameters
    ----------
        form_data: dict
            данные формы
        plan_name: str
            название учедного плана
        wp_data: dict
            данные кабочих программ

    Returns
    -------
        str
            название файла
    """
    document = Document()
    section = document.sections[-1]
    # Отступы
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    paragraph_format = document.styles["Normal"].paragraph_format
    # межстрочный интервал
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # между абзацами
    paragraph_format.space_after = Pt(0)
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    for wp in wp_data.values():
        wp_name = wp.get("name")

        title = document.add_paragraph("")
        title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        line1 = title.add_run("МИНИСТЕРСТВО ВНУТРЕННИХ ДЕЛ РОССИЙСКОЙ ФЕДЕРАЦИИ\n")
        line1.font.size = Pt(13)
        line1.bold = True
        line2 = title.add_run(
            form_data.get("organization_name").replace("\r", "").replace(" «", "\n«")
        )
        line2.font.size = Pt(13)
        line2.bold = True

        document.add_paragraph("")
        document.add_paragraph("")

        approval_info = document.add_paragraph(
            form_data.get("wp_approval_info").replace("\r", "")
        )
        approval_info.paragraph_format.left_indent = Cm(9)
        approval_info.add_run(f"\n{form_data.get('chief_rank')}\n")

        chief_name = document.add_paragraph(form_data.get("chief_name"))
        chief_name.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        date_info = document.add_paragraph(f"« {form_data.get('wp_approval_day')} »")
        date_info.paragraph_format.left_indent = Cm(9)
        date_info.add_run(
            f" {form_data.get('wp_approval_month').center(20)} "
        ).underline = True
        date_info.add_run(f"{form_data.get('wp_approval_year')} г.")

        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("")

        program_info = document.add_paragraph(f"{form_data.get('wp_type_name')}\n")
        program_info.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        program_info.add_run(f"«{wp_name}»\n").bold = True
        program_info.add_run(f"{form_data.get('wp_speciality_type')}\n")
        program_info.add_run(f"{form_data.get('wp_speciality')}\n")
        program_info.add_run(
            f"{form_data.get('wp_specialization_type')} - "
            f"{form_data.get('wp_specialization')}\n"
        ).italic = True
        if form_data.get("switch_narrow_spec"):
            program_info.add_run(
                f"узкая специализация - {form_data.get('wp_narrow_specialization')}\n"
            ).italic = True
        if form_data.get("switch_foreign"):
            program_info.add_run(f"иностранные слушатели\n").italic = True
        program_info.add_run(
            f"форма обучения - {form_data.get('wp_education_form')}\n"
        ).italic = True
        program_info.add_run(f"набор {form_data.get('wp_year')} года\n\n").italic = True
        if form_data.get("switch_qualification"):
            program_info.add_run(f"квалификация - {form_data.get('wp_qualification')}")

        document.add_page_break()

    filename = f"{plan_name}.docx"
    document.save(FlaskConfig.EXPORT_FILE_DIR + filename)
    logging.debug(f"Сформирован файл - титульные листы РП: {filename}")
    return filename
