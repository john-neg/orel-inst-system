import logging
import os

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt

from config import FlaskConfig, ApeksConfig as Apeks


def generate_indicators_file(plan_name: str, report_data: dict) -> str:
    """Формирует отчет об индикаторах по дисциплинам в формате docx.

    :param plan_name:str: название учебного плана
    :param report_data:dict: данные об индикаторах по дисциплинам
        {'disc_name': {'comp_list': [values], 'indicator_type': [values]}}
    :return: Название сформированного файла
    """

    document = Document()
    section = document.sections[-1]
    section.top_margin = Cm(1)  # Верхний отступ
    section.bottom_margin = Cm(1)  # Нижний отступ
    section.left_margin = Cm(1)  # Отступ слева
    section.right_margin = Cm(1)  # Отступ справа
    paragraph_format = document.styles["Normal"].paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # межстрочный интервал
    paragraph_format.space_after = Pt(0)  # между абзацами
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"  # Стиль шрифта
    font.size = Pt(12)
    document.add_heading(plan_name, 0)  # Заголовок
    document.add_paragraph("Индикаторы компетенций по дисциплинам")
    for disc in report_data:
        document.add_page_break()
        document.add_heading(disc, level=1)
        # Список компетенций
        comp_list = report_data[disc].get("comp_list")
        document.add_paragraph(", ".join(comp_list))
        # Индикаторы
        for typ in Apeks.INDICATOR_TYPES:
            document.add_paragraph("")
            p = document.add_paragraph("")
            p.add_run(f"{Apeks.INDICATOR_TYPES[typ]}:").bold = True
            for ind in report_data[disc].get(typ):
                paragraph = document.add_paragraph(f"{ind};")
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    filename = f"{plan_name}.docx"
    document.save(os.path.join(FlaskConfig.EXPORT_FILE_DIR, filename))
    logging.debug(f"Сформирован файл - отчет по индикаторам: {filename}")
    return filename
